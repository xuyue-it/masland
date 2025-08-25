from flask import Flask, render_template, request, redirect, url_for, send_file, jsonify, session
from docx import Document
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.header import Header
from email.utils import formataddr
import traceback
from dotenv import load_dotenv
load_dotenv()
import os
from functools import wraps
import psycopg2

app = Flask(__name__)

# ========== åŸºæœ¬é…ç½® ==========
app.secret_key = os.getenv("SECRET_KEY", "replace-this-in-prod")

# ç®¡ç†å‘˜å¯†ç 
ADMIN_PASSWORD = "maslandit339188"
print(">>> ADMIN_PASSWORD source: CODE")

# ========== é‚®ä»¶é…ç½® ==========
SMTP_SERVER = os.getenv("SMTP_SERVER", "smtp.gmail.com")
SMTP_PORT   = int(os.getenv("SMTP_PORT", "587"))
SENDER_EMAIL    = os.getenv("SENDER_EMAIL", "jameslau32@gmail.com")
SENDER_PASSWORD = os.getenv("SENDER_PASSWORD", "nkxp hwba oagf vgyq")
ADMIN_EMAIL     = os.getenv("ADMIN_EMAIL", "jameslau32@gmail.com")

# ========== æ•°æ®åº“é…ç½®ï¼ˆNeon PostgreSQLï¼‰ ==========
DB_URL = os.getenv("DATABASE_URL")

def get_conn():
    return psycopg2.connect(DB_URL)

# ========== å™¨ææ¸…å• ==========
EQUIP_MAP = {
    "mic": "éº¦å…‹é£","amp": "æ‰©éŸ³å™¨","pa": "éŸ³å“ç³»ç»Ÿ","projector": "æŠ•å½±æœº","screen": "æŠ•å½±å±å¹•",
    "ext": "å»¶é•¿çº¿","table": "æ¡Œå­","chair": "æ¤…å­","podium": "è®²å°","hdmi": "HDMIçº¿",
}

# ===== é‚®ä»¶å‘é€ =====
def send_email(subject, content, to_email):
    msg = MIMEMultipart()
    msg['From'] = formataddr(("ç¦æºå ‚å™¨æå¤–å€Ÿç³»ç»Ÿ", SENDER_EMAIL))
    msg['To'] = to_email
    msg['Subject'] = Header(subject, "utf-8")
    msg.attach(MIMEText(content, "plain", "utf-8"))
    try:
        server = smtplib.SMTP_SSL(SMTP_SERVER, 465, timeout=20)
        server.login(SENDER_EMAIL, SENDER_PASSWORD)
        server.sendmail(SENDER_EMAIL, [to_email], msg.as_string())
        server.quit()
        print(f"âœ… é‚®ä»¶å·²å‘é€è‡³ {to_email}ï¼ˆSSL:465ï¼‰")
        return True, None
    except Exception as e_ssl:
        print("âš ï¸ SSL(465) å‘é€å¤±è´¥ï¼š", e_ssl)
        print(traceback.format_exc())
    try:
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT, timeout=20)
        server.ehlo(); server.starttls(); server.ehlo()
        server.login(SENDER_EMAIL, SENDER_PASSWORD)
        server.sendmail(SENDER_EMAIL, [to_email], msg.as_string())
        server.quit()
        print(f"âœ… é‚®ä»¶å·²å‘é€è‡³ {to_email}ï¼ˆTLS:587ï¼‰")
        return True, None
    except Exception as e_tls:
        print("âŒ é‚®ä»¶å‘é€å¤±è´¥ï¼ˆTLS:587ï¼‰ï¼š", e_tls)
        print(traceback.format_exc())
        return False, str(e_tls)

# ========================
# æ•°æ®åº“åˆå§‹åŒ–
# ========================
def init_db():
    conn = get_conn()
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS submissions (
        id SERIAL PRIMARY KEY,
        name TEXT, phone TEXT, email TEXT, group_name TEXT, event_name TEXT,
        start_date TEXT, start_time TEXT, end_date TEXT, end_time TEXT,
        location TEXT, event_type TEXT, participants TEXT, equipment TEXT,
        special_request TEXT, donation TEXT, donation_method TEXT,
        remarks TEXT, emergency_name TEXT, emergency_phone TEXT,
        status TEXT DEFAULT 'å¾…å®¡æ ¸', review_comment TEXT
    )''')
    conn.commit(); conn.close()

init_db()

# ========================
# ç™»å½•ä¿æŠ¤
# ========================
def login_required(view_func):
    @wraps(view_func)
    def wrapper(*args, **kwargs):
        if not session.get("logged_in"):
            return redirect(url_for("login", next=request.path))
        return view_func(*args, **kwargs)
    return wrapper

@app.route("/login", methods=["GET", "POST"])
def login():
    error = None
    if request.method == "POST":
        if request.form.get("password", "") == ADMIN_PASSWORD:
            session["logged_in"] = True
            return redirect(request.args.get("next") or url_for("admin"))
        error = "å¯†ç é”™è¯¯ï¼Œè¯·é‡è¯•ã€‚"
    return render_template("login.html", error=error)

@app.route("/logout")
def logout():
    session.pop("logged_in", None)
    return redirect(url_for("login"))

# ========================
# å‰å°
# ========================
@app.route("/")
def index():
    return render_template("index.html", equip_map=EQUIP_MAP)

@app.route("/submit", methods=["POST"])
def submit():
    data = request.form.to_dict(flat=True)
    equip_items = []
    for key, cname in EQUIP_MAP.items():
        if data.get(f"equip_{key}") == "on":
            qty_str = (data.get(f"equip_{key}_qty") or "").strip()
            try: qty = int(qty_str)
            except: qty = 0
            if qty <= 0: qty = 1
            equip_items.append(f"{cname}x{qty}")
    equipment_str = ", ".join(equip_items)

    conn = get_conn(); c = conn.cursor()
    c.execute('''
        INSERT INTO submissions (
            name, phone, email, group_name, event_name,
            start_date, start_time, end_date, end_time,
            location, event_type, participants, equipment,
            special_request, donation, donation_method,
            remarks, emergency_name, emergency_phone
        ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
    ''', (
        data.get('name'), data.get('phone'), data.get('email'), data.get('group'),
        data.get('event_name'), data.get('start_date'), data.get('start_time'),
        data.get('end_date'), data.get('end_time'), data.get('location'),
        data.get('event_type'), data.get('participants'), equipment_str,
        data.get('special_request'), data.get('donation'), data.get('donation_method'),
        data.get('remarks'), data.get('emergency_name'), data.get('emergency_phone')
    ))
    conn.commit(); conn.close()

    send_email("ã€æ–°ç”³è¯·ã€‘ç¦æºå ‚å™¨æå¤–å€Ÿ",
               f"ç”³è¯·äººï¼š{data.get('name')}\næ´»åŠ¨ï¼š{data.get('event_name')}\nç”µè¯ï¼š{data.get('phone')}\né‚®ç®±ï¼š{data.get('email')}",
               ADMIN_EMAIL)

    return """<html><body><h1>æäº¤æˆåŠŸï¼</h1><p>è¯·è¿”å›é¦–é¡µæŸ¥è¯¢å®¡æ ¸çŠ¶æ€ã€‚</p></body></html>"""

# ========================
# ç®¡ç†é¡µ + æ¥å£
# ========================
@app.route("/admin")
@login_required
def admin():
    conn = get_conn(); c = conn.cursor()
    c.execute("SELECT * FROM submissions ORDER BY id DESC")
    submissions = c.fetchall()
    conn.close()
    return render_template("admin.html", submissions=submissions)

@app.route("/api/submission/<int:submission_id>")
@login_required
def api_submission(submission_id):
    conn = get_conn(); c = conn.cursor()
    c.execute("""SELECT id, name, email, event_name, status, review_comment
                 FROM submissions WHERE id=%s""", (submission_id,))
    row = c.fetchone(); conn.close()
    if not row:
        return jsonify({"success": False, "message": "è®°å½•ä¸å­˜åœ¨"}), 404
    return jsonify({"success": True,"data": {
        "id": row[0], "name": row[1], "email": row[2], "event_name": row[3],
        "status": row[4] or "å¾…å®¡æ ¸", "review_comment": row[5] or ""
    }})

@app.route("/update_status/<int:submission_id>/<string:new_status>", methods=["POST"])
@login_required
def update_status(submission_id, new_status):
    try:
        data = request.get_json(silent=True) or {}
        comment = data.get("comment", "")

        conn = get_conn(); c = conn.cursor()
        c.execute("UPDATE submissions SET status=%s, review_comment=%s WHERE id=%s",
                  (new_status, comment, submission_id))
        conn.commit()

        # å–å›æœ€æ–°æ•°æ®
        c.execute("SELECT name, email, status FROM submissions WHERE id=%s", (submission_id,))
        row = c.fetchone()
        conn.close()

        if row and row[1]:  # é‚®ä»¶é€šçŸ¥
            try:
                send_email("ã€å®¡æ ¸ç»“æœã€‘ç¦æºå ‚å™¨æå¤–å€Ÿç”³è¯·",
                           f"æ‚¨å¥½ {row[0]}ï¼Œæ‚¨çš„ç”³è¯·å·²è¢«å®¡æ ¸ä¸ºï¼š{row[2]}\nå®¡æ ¸è¯´æ˜ï¼š{comment or 'æ— '}",
                           row[1])
            except Exception as mail_err:
                print("âš ï¸ å®¡æ ¸åé€šçŸ¥ç”³è¯·äººå¤±è´¥ï¼ˆå¿½ç•¥ï¼‰ï¼š", mail_err)

        return jsonify({"success": True, "submission_id": submission_id,
                        "name": row[0] if row else "", "status": row[2] if row else new_status})
    except Exception as e:
        print("âŒ /update_status å‡ºé”™ï¼š", e); print(traceback.format_exc())
        return jsonify({"success": False, "message": f"æœåŠ¡å™¨é”™è¯¯ï¼š{e}"}), 500

@app.route("/send_review_email/<int:submission_id>", methods=["POST"])
@login_required
def send_review_email(submission_id):
    conn = get_conn(); c = conn.cursor()
    c.execute("SELECT name, email, event_name, status, review_comment FROM submissions WHERE id=%s",
              (submission_id,))
    row = c.fetchone(); conn.close()
    if not row: return jsonify({"success": False, "message": "è®°å½•ä¸å­˜åœ¨"}), 404

    name, email, event_name, status, review_comment = row
    if not email: return jsonify({"success": False, "message": "è¯¥è®°å½•æ²¡æœ‰å¡«å†™é‚®ç®±ï¼Œæ— æ³•å‘é€"}), 400

    ok, err = send_email(f"ã€å®¡æ ¸ç»“æœã€‘{event_name or ''}",
                         f"æ‚¨å¥½ {name or ''}ï¼š\n\næ‚¨çš„ç”³è¯·ï¼ˆæ´»åŠ¨ï¼š{event_name or '-'}) "
                         f"å®¡æ ¸ç»“æœä¸ºï¼š{status or 'å¾…å®¡æ ¸'}\nå®¡æ ¸è¯´æ˜ï¼š{review_comment or 'æ— '}\n\n"
                         f"å¦‚æœ‰ç–‘é—®è¯·å›å¤æ­¤é‚®ä»¶è”ç³»ç®¡ç†å‘˜ã€‚",
                         email)
    if ok: return jsonify({"success": True, "message": f"å·²å‘é€åˆ° {email}"})
    else:  return jsonify({"success": False, "message": f"å‘é€å¤±è´¥ï¼š{err}"}), 500

@app.route("/delete_submission/<int:submission_id>", methods=["POST"])
@login_required
def delete_submission(submission_id):
    conn = get_conn(); c = conn.cursor()
    c.execute("DELETE FROM submissions WHERE id=%s", (submission_id,))
    affected = c.rowcount
    conn.commit(); conn.close()
    return jsonify({"success": True, "submission_id": submission_id, "deleted": affected})

@app.route("/download/<int:submission_id>")
@login_required
def download(submission_id):
    conn = get_conn(); c = conn.cursor()
    c.execute("SELECT * FROM submissions WHERE id=%s", (submission_id,))
    submission = c.fetchone(); conn.close()
    if not submission: return "è®°å½•ä¸å­˜åœ¨"

    doc = Document(); doc.add_heading('ç”³è¯·è¡¨è¯¦æƒ…', level=1)
    fields = ["ID","å§“å","ç”µè¯","é‚®ç®±","å›¢ä½“åç§°","æ´»åŠ¨åç§°","å¼€å§‹æ—¥æœŸ","å¼€å§‹æ—¶é—´",
              "ç»“æŸæ—¥æœŸ","ç»“æŸæ—¶é—´","åœ°ç‚¹","æ´»åŠ¨ç±»å‹","å‚ä¸äººæ•°","å™¨æ","ç‰¹åˆ«éœ€æ±‚",
              "ææ¬¾","ææ¬¾æ–¹å¼","å¤‡æ³¨","ç´§æ€¥è”ç³»äºº","ç´§æ€¥è”ç³»ç”µè¯","å®¡æ ¸çŠ¶æ€","å®¡æ ¸è¯´æ˜"]
    for i, field in enumerate(fields):
        if i < len(submission):
            doc.add_paragraph(f"{field}: {submission[i]}")

    file_path = f"submission_{submission_id}.docx"; doc.save(file_path)
    return send_file(file_path, as_attachment=True)

# ========================
# ğŸ” æŸ¥è¯¢çŠ¶æ€ APIï¼ˆæ¢å¤ï¼‰
# ========================
@app.route("/check_status_api")
def check_status_api():
    name = request.args.get("name")
    if not name:
        return jsonify({"status": "error", "message": "Name is required"})

    conn = get_conn(); c = conn.cursor()
    c.execute("SELECT name, event_name, status, review_comment FROM submissions WHERE name=%s ORDER BY id DESC LIMIT 1", (name,))
    row = c.fetchone(); conn.close()

    if row:
        return jsonify({
            "status": row[2],
            "data": {
                "name": row[0],
                "event_name": row[1],
                "review_status": row[2],
                "review_comment": row[3] or ""
            }
        })
    else:
        return jsonify({"status": "not_found"})

@app.route("/_health")
def _health(): return "ok", 200

if __name__ == "__main__":
    app.run(debug=True)
